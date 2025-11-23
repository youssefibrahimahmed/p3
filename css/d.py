from docx import Document

# إنشاء المستند
doc = Document()

# إضافة العنوان الرئيسي
title = doc.add_heading("Recent Scandals for IT Companies", level=1)
title.alignment = 1  # توسيط العنوان

# معلومات الطالب
doc.add_paragraph("Name: Youssef Ibrahim Ahmed")
doc.add_paragraph("Group: 11")
doc.add_paragraph("Year: Second Year")
doc.add_paragraph("")

# المقدمة
doc.add_heading("Introduction", level=2)
doc.add_paragraph(
    "In recent years, information technology (IT) companies have become some of the most powerful "
    "and influential organizations in the world. Their products and services shape how people communicate, "
    "work, and live. However, with great influence comes great responsibility — and sometimes controversy. "
    "Several IT companies have faced serious scandals involving data privacy violations, unethical business practices, "
    "or cybersecurity failures. This research aims to highlight some of the most recent scandals that affected major "
    "technology firms and to understand their impact on public trust and corporate reputation."
)

# Huawei
doc.add_heading("1. Huawei – Allegations of Bribery and Political Influence", level=2)
doc.add_paragraph(
    "In 2025, Huawei, one of the world’s largest telecommunications companies, faced a major investigation "
    "in Belgium and other parts of Europe. Authorities accused the company of attempting to bribe members of "
    "the European Parliament to influence decisions related to 5G network regulations. This scandal reignited debates "
    "about the relationship between technology companies and governments, as well as national security concerns "
    "regarding Chinese technology in Europe."
)

# Meta
doc.add_heading("2. Meta (Facebook) – Ongoing Privacy and Data Misuse Issues", level=2)
doc.add_paragraph(
    "Meta has been repeatedly accused of mishandling user data. Despite the Cambridge Analytica scandal in 2018, "
    "new reports in 2024 and 2025 showed that Meta continued to collect excessive personal data from users through "
    "its apps, including Instagram and WhatsApp. Privacy regulators in the EU and the United States fined the company "
    "billions of dollars for breaching data protection laws. These incidents have damaged the company’s reputation and "
    "sparked global discussions about the ethics of social media data collection."
)

# OpenAI and Google
doc.add_heading("3. OpenAI and Google – Concerns over AI Transparency", level=2)
doc.add_paragraph(
    "With the rapid development of artificial intelligence, companies like OpenAI and Google have been criticized "
    "for their lack of transparency in how they train large language models. Reports in 2025 revealed that some AI "
    "datasets included copyrighted materials and personal data without proper consent. This raised ethical and legal "
    "questions about intellectual property, data protection, and the fairness of AI systems."
)

# Big Technologies
doc.add_heading("4. Big Technologies PLC – Financial Misconduct", level=2)
doc.add_paragraph(
    "In early 2025, Big Technologies, a UK-based tech company, dismissed its CEO after discovering financial irregularities "
    "and undeclared offshore transactions. The case illustrated how corporate governance issues can affect even highly successful "
    "IT firms, leading to loss of investor confidence and reputational harm."
)

# الخاتمة
doc.add_heading("Conclusion", level=2)
doc.add_paragraph(
    "These scandals show that even the most advanced IT companies are not immune to ethical, legal, and social challenges. "
    "Whether it involves privacy, transparency, or corporate integrity, each scandal highlights the importance of accountability "
    "in the technology sector. To maintain public trust, IT companies must adopt stronger ethical standards, ensure compliance "
    "with global regulations, and prioritize transparency in their operations."
)

# حفظ الملف
doc.save("Recent_Scandals_for_IT_Companies_Youssef_Ibrahim_Group11.docx")
